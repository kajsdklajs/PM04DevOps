import random
import os
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from tkinter.simpledialog import Dialog
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.chart import LineChart, Reference
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class MeasurementGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор файлов измерений")
        self.root.geometry("1000x800")
        
        self.measurements = []
        self.setup_ui()
        
    def setup_ui(self):
        # Основной фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Заголовок
        title_label = ttk.Label(main_frame, text="ГЕНЕРАТОР ФАЙЛОВ ИЗМЕРЕНИЙ", 
                               font=('Arial', 16, 'bold'))
        title_label.pack(pady=(0, 20))
        
        # Фрейм параметров
        params_frame = ttk.LabelFrame(main_frame, text="Параметры измерений", padding="10")
        params_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Параметры ввода
        input_frame = ttk.Frame(params_frame)
        input_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(input_frame, text="Количество измерений:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.sample_size_var = tk.StringVar(value="10")
        ttk.Entry(input_frame, textvariable=self.sample_size_var, width=10).grid(row=0, column=1, padx=(0, 30))
        
        ttk.Label(input_frame, text="Способ заполнения:").grid(row=0, column=2, sticky=tk.W, padx=(0, 10))
        self.fill_method_var = tk.StringVar(value="auto")
        
        ttk.Radiobutton(input_frame, text="Автоматически", variable=self.fill_method_var, 
                       value="auto", command=self.toggle_input_method).grid(row=0, column=3, padx=(0, 20))
        ttk.Radiobutton(input_frame, text="Вручную", variable=self.fill_method_var, 
                       value="manual", command=self.toggle_input_method).grid(row=0, column=4)
        
        # Фрейм для ручного ввода
        self.manual_input_frame = ttk.LabelFrame(params_frame, text="Ручной ввод значений", padding="10")
        self.manual_input_frame.pack(fill=tk.X, pady=5)
        
        self.manual_input_text = scrolledtext.ScrolledText(self.manual_input_frame, height=4, width=80)
        self.manual_input_text.pack(fill=tk.BOTH, expand=True)
        ttk.Label(self.manual_input_frame, text="Введите значения через пробел или с новой строки (целые числа 0-100):", 
                 font=('Arial', 8)).pack(pady=(5, 0))
        
        # Скрыть ручной ввод по умолчанию
        self.manual_input_frame.pack_forget()
        
        # Фрейм кнопок управления
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(pady=10)
        
        ttk.Button(buttons_frame, text="Сгенерировать данные", 
                  command=self.generate_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Создать Word документ", 
                  command=self.create_word_document).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Создать Excel файл", 
                  command=self.create_excel_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Создать оба файла", 
                  command=self.create_both_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Очистить", 
                  command=self.clear_all).pack(side=tk.LEFT, padx=5)
        
        # Фрейм для отображения данных
        display_frame = ttk.LabelFrame(main_frame, text="Результаты", padding="10")
        display_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        
        # Вкладки
        notebook = ttk.Notebook(display_frame)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # Вкладка данных
        data_tab = ttk.Frame(notebook)
        notebook.add(data_tab, text="Таблица измерений")
        self.data_text = scrolledtext.ScrolledText(data_tab, width=80, height=12, font=('Courier', 9))
        self.data_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Вкладка графика
        plot_tab = ttk.Frame(notebook)
        notebook.add(plot_tab, text="График измерений")
        self.setup_plot_tab(plot_tab)
        
        # Вкладка статистики
        stats_tab = ttk.Frame(notebook)
        notebook.add(stats_tab, text="Статистика")
        self.stats_text = scrolledtext.ScrolledText(stats_tab, width=80, height=12)
        self.stats_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Статус бар
        self.status_var = tk.StringVar(value="Готов к работе")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        status_bar.pack(fill=tk.X, pady=(10, 0))
        
        # Проверка доступности библиотек
        self.check_dependencies()
    
    def setup_plot_tab(self, parent):
        """Настройка вкладки с графиком"""
        # Фрейм для управления графиком
        plot_controls_frame = ttk.Frame(parent)
        plot_controls_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(plot_controls_frame, text="Обновить график", 
                  command=self.update_plot).pack(side=tk.LEFT, padx=5)
        
        # Фрейм для графика
        plot_frame = ttk.Frame(parent)
        plot_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Создание фигуры matplotlib
        self.fig = Figure(figsize=(8, 4), dpi=100)
        self.ax = self.fig.add_subplot(111)
        
        # Холст для графика
        self.canvas = FigureCanvasTkAgg(self.fig, plot_frame)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        
        # Начальный пустой график
        self.ax.set_title("График результатов измерений")
        self.ax.set_xlabel("Номер измерения")
        self.ax.set_ylabel("Результат")
        self.ax.grid(True, alpha=0.3)
        self.ax.set_ylim(0, 100)
        self.canvas.draw()
    
    def check_dependencies(self):
        """Проверка доступности необходимых библиотек"""
        missing_libs = []
        if not DOCX_AVAILABLE:
            missing_libs.append("python-docx")
        if not EXCEL_AVAILABLE:
            missing_libs.append("openpyxl")
        
        if missing_libs:
            messagebox.showwarning(
                "Внимание", 
                f"Не установлены библиотеки: {', '.join(missing_libs)}\n\n"
                f"Установите их с помощью:\n"
                f"pip install {' '.join(missing_libs)}"
            )
    
    def toggle_input_method(self):
        """Переключение между методами ввода"""
        if self.fill_method_var.get() == "manual":
            self.manual_input_frame.pack(fill=tk.X, pady=5)
        else:
            self.manual_input_frame.pack_forget()
    
    def generate_data(self):
        """Генерация данных измерений"""
        try:
            self.status_var.set("Генерация данных...")
            self.root.update()
            
            sample_size = int(self.sample_size_var.get())
            if sample_size <= 0:
                messagebox.showerror("Ошибка", "Количество измерений должно быть положительным")
                return
            
            if self.fill_method_var.get() == "auto":
                # Автоматическая генерация
                self.measurements = [random.randint(0, 100) for _ in range(sample_size)]
            else:
                # Ручной ввод
                input_text = self.manual_input_text.get(1.0, tk.END).strip()
                if not input_text:
                    messagebox.showerror("Ошибка", "Введите значения измерений")
                    return
                
                # Разбор введенных значений
                values = []
                for line in input_text.split('\n'):
                    for value in line.split():
                        try:
                            num = int(value.strip())
                            if 0 <= num <= 100:
                                values.append(num)
                            else:
                                messagebox.showerror("Ошибка", f"Значение {num} вне диапазона 0-100")
                                return
                        except ValueError:
                            messagebox.showerror("Ошибка", f"Некорректное значение: {value}")
                            return
                
                if len(values) != sample_size:
                    messagebox.showerror("Ошибка", 
                                       f"Введено {len(values)} значений, ожидается {sample_size}")
                    return
                
                self.measurements = values
            
            # Отображение данных
            self.display_data()
            self.update_plot()
            self.calculate_statistics()
            
            self.status_var.set(f"Данные сгенерированы: {sample_size} измерений")
            
        except ValueError as e:
            messagebox.showerror("Ошибка ввода", "Проверьте правильность введенных значений")
            self.status_var.set("Ошибка ввода данных")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")
            self.status_var.set("Ошибка при генерации данных")
    
    def display_data(self):
        """Отображение данных в таблице"""
        self.data_text.delete(1.0, tk.END)
        
        # Заголовок таблицы
        self.data_text.insert(tk.END, f"{'№':<4} {'Результат':<10}\n")
        self.data_text.insert(tk.END, "-" * 15 + "\n")
        
        # Данные таблицы
        for i, measurement in enumerate(self.measurements, 1):
            self.data_text.insert(tk.END, f"{i:<4} {measurement:<10}\n")
    
    def update_plot(self):
        """Обновление графика"""
        if not self.measurements:
            return
        
        self.ax.clear()
        
        # Построение графика
        x = list(range(1, len(self.measurements) + 1))
        self.ax.plot(x, self.measurements, 'o-', linewidth=2, markersize=4, color='blue', alpha=0.7)
        
        # Настройки графика
        self.ax.set_title("График результатов измерений", fontsize=12)
        self.ax.set_xlabel("Номер измерения", fontsize=10)
        self.ax.set_ylabel("Результат", fontsize=10)
        self.ax.grid(True, alpha=0.3)
        self.ax.set_ylim(0, 100)
        self.ax.set_xlim(0.5, len(self.measurements) + 0.5)
        
        # Добавление горизонтальных линий для среднего
        avg_value = sum(self.measurements) / len(self.measurements)
        self.ax.axhline(y=avg_value, color='red', linestyle='--', alpha=0.7, 
                       label=f'Среднее: {avg_value:.1f}')
        
        self.ax.legend()
        self.canvas.draw()
    
    def calculate_statistics(self):
        """Расчет статистики"""
        if not self.measurements:
            return
        
        self.stats_text.delete(1.0, tk.END)
        
        stats = {
            'Объем выборки': len(self.measurements),
            'Минимальное значение': min(self.measurements),
            'Максимальное значение': max(self.measurements),
            'Среднее значение': sum(self.measurements) / len(self.measurements),
            'Сумма': sum(self.measurements),
            'Медиана': self.calculate_median(),
            'Стандартное отклонение': self.calculate_std_dev(),
            'Количество значений > 50': sum(1 for x in self.measurements if x > 50),
            'Количество значений < 50': sum(1 for x in self.measurements if x < 50),
            'Количество значений = 50': sum(1 for x in self.measurements if x == 50)
        }
        
        self.stats_text.insert(tk.END, "СТАТИСТИЧЕСКИЕ ПОКАЗАТЕЛИ\n")
        self.stats_text.insert(tk.END, "=" * 40 + "\n\n")
        
        for key, value in stats.items():
            if isinstance(value, float):
                self.stats_text.insert(tk.END, f"{key}: {value:.2f}\n")
            else:
                self.stats_text.insert(tk.END, f"{key}: {value}\n")
    
    def calculate_median(self):
        """Расчет медианы"""
        sorted_measurements = sorted(self.measurements)
        n = len(sorted_measurements)
        if n % 2 == 0:
            return (sorted_measurements[n//2 - 1] + sorted_measurements[n//2]) / 2
        else:
            return sorted_measurements[n//2]
    
    def calculate_std_dev(self):
        """Расчет стандартного отклонения"""
        mean = sum(self.measurements) / len(self.measurements)
        variance = sum((x - mean) ** 2 for x in self.measurements) / len(self.measurements)
        return variance ** 0.5
    
    def create_word_document(self):
        """Создание Word документа"""
        if not self.measurements:
            messagebox.showwarning("Предупреждение", "Сначала сгенерируйте данные")
            return
        
        if not DOCX_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека python-docx не установлена")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить Word документ",
            initialfile=f"measurements_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        if filename:
            try:
                self._create_word_file(filename)
                self.status_var.set(f"Word документ создан: {os.path.basename(filename)}")
                messagebox.showinfo("Успех", f"Word документ сохранен:\n{filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось создать Word документ: {e}")
    
    def _create_word_file(self, filename):
        """Создание Word файла"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('Результаты измерений', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о данных
        doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        doc.add_paragraph(f"Количество измерений: {len(self.measurements)}")
        doc.add_paragraph(f"Способ заполнения: {'Автоматически' if self.fill_method_var.get() == 'auto' else 'Вручную'}")
        doc.add_paragraph()
        
        # Создание таблицы
        table = doc.add_table(rows=len(self.measurements) + 1, cols=2)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Номер измерения'
        header_cells[1].text = 'Результат'
        
        # Форматирование заголовков
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Заполнение таблицы данными
        for i, measurement in enumerate(self.measurements, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(measurement)
            
            # Выравнивание по центру
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавление статистики
        doc.add_page_break()
        stats_heading = doc.add_heading('Статистика измерений', level=1)
        
        stats_text = self.stats_text.get(1.0, tk.END)
        for line in stats_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(filename)
    
    def create_excel_file(self):
        """Создание Excel файла"""
        if not self.measurements:
            messagebox.showwarning("Предупреждение", "Сначала сгенерируйте данные")
            return
        
        if not EXCEL_AVAILABLE:
            messagebox.showerror("Ошибка", "Библиотека openpyxl не установлена")
            return
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Все файлы", "*.*")],
            title="Сохранить Excel файл",
            initialfile=f"measurements_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        )
        
        if filename:
            try:
                self._create_excel_file(filename)
                self.status_var.set(f"Excel файл создан: {os.path.basename(filename)}")
                messagebox.showinfo("Успех", f"Excel файл сохранен:\n{filename}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось создать Excel файл: {e}")
    
    def _create_excel_file(self, filename):
        """Создание Excel файла"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Результаты измерений"
        
        # Настройка стилей
        header_font = Font(bold=True, size=12)
        border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                       top=Side(style='thin'), bottom=Side(style='thin'))
        center_align = Alignment(horizontal='center', vertical='center')
        
        # Заголовок
        ws.merge_cells('A1:B1')
        ws['A1'] = 'Результаты измерений'
        ws['A1'].font = Font(bold=True, size=14)
        ws['A1'].alignment = center_align
        
        # Информация о данных
        ws['A3'] = 'Дата создания:'
        ws['B3'] = datetime.now().strftime('%d.%m.%Y %H:%M:%S')
        ws['A4'] = 'Количество измерений:'
        ws['B4'] = len(self.measurements)
        ws['A5'] = 'Способ заполнения:'
        ws['B5'] = 'Автоматически' if self.fill_method_var.get() == 'auto' else 'Вручную'
        
        # Заголовки таблицы
        ws['A7'] = 'Номер измерения'
        ws['B7'] = 'Результат'
        
        # Применение стилей к заголовкам
        for cell in ['A7', 'B7']:
            ws[cell].font = header_font
            ws[cell].border = border
            ws[cell].alignment = center_align
        
        # Заполнение данными
        for i, measurement in enumerate(self.measurements, 1):
            ws[f'A{7 + i}'] = i
            ws[f'B{7 + i}'] = measurement
            
            # Применение стилей к ячейкам данных
            for cell in [f'A{7 + i}', f'B{7 + i}']:
                ws[cell].border = border
                ws[cell].alignment = center_align
        
        # Автоподбор ширины столбцов
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 15
        
        # Создание листа со статистикой
        ws_stats = wb.create_sheet("Статистика")
        
        stats_text = self.stats_text.get(1.0, tk.END)
        lines = [line for line in stats_text.split('\n') if line.strip()]
        
        ws_stats['A1'] = 'Статистический показатель'
        ws_stats['B1'] = 'Значение'
        ws_stats['A1'].font = header_font
        ws_stats['B1'].font = header_font
        
        for i, line in enumerate(lines, 2):
            if ':' in line:
                key, value = line.split(':', 1)
                ws_stats[f'A{i}'] = key.strip()
                ws_stats[f'B{i}'] = value.strip()
        
        # Автоподбор ширины для листа статистики
        ws_stats.column_dimensions['A'].width = 25
        ws_stats.column_dimensions['B'].width = 15
        
        # Создание диаграммы
        self._create_excel_chart(ws, len(self.measurements))
        
        wb.save(filename)
    
    def _create_excel_chart(self, worksheet, data_count):
        """Создание диаграммы в Excel"""
        try:
            chart = LineChart()
            chart.title = "График результатов измерений"
            chart.style = 13
            chart.x_axis.title = "Номер измерения"
            chart.y_axis.title = "Результат"
            
            data = Reference(worksheet, min_col=2, min_row=7, max_row=7 + data_count)
            categories = Reference(worksheet, min_col=1, min_row=8, max_row=7 + data_count)
            
            chart.add_data(data, titles_from_data=True)
            chart.set_categories(categories)
            
            # Размещение диаграммы на листе
            worksheet.add_chart(chart, "D2")
            
        except Exception as e:
            print(f"Диаграмма не создана: {e}")
    
    def create_both_files(self):
        """Создание обоих файлов"""
        if not self.measurements:
            messagebox.showwarning("Предупреждение", "Сначала сгенерируйте данные")
            return
        
        if not DOCX_AVAILABLE or not EXCEL_AVAILABLE:
            messagebox.showerror("Ошибка", "Не установлены необходимые библиотеки")
            return
        
        base_filename = filedialog.asksaveasfilename(
            title="Сохранить файлы измерений",
            initialfile=f"measurements_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            filetypes=[("All supported", "*.docx;*.xlsx"), ("Все файлы", "*.*")]
        )
        
        if base_filename:
            try:
                # Убираем расширение если есть
                base_name = os.path.splitext(base_filename)[0]
                
                word_filename = base_name + ".docx"
                excel_filename = base_name + ".xlsx"
                
                self._create_word_file(word_filename)
                self._create_excel_file(excel_filename)
                
                self.status_var.set("Оба файла созданы")
                messagebox.showinfo("Успех", 
                                  f"Файлы успешно созданы:\n"
                                  f"Word: {word_filename}\n"
                                  f"Excel: {excel_filename}")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось создать файлы: {e}")
    
    def clear_all(self):
        """Очистка всех данных"""
        self.measurements = []
        self.data_text.delete(1.0, tk.END)
        self.stats_text.delete(1.0, tk.END)
        self.manual_input_text.delete(1.0, tk.END)
        
        # Очистка графика
        self.ax.clear()
        self.ax.set_title("График результатов измерений")
        self.ax.set_xlabel("Номер измерения")
        self.ax.set_ylabel("Результат")
        self.ax.grid(True, alpha=0.3)
        self.ax.set_ylim(0, 100)
        self.canvas.draw()
        
        self.status_var.set("Данные очищены")

def main():
    root = tk.Tk()
    app = MeasurementGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()